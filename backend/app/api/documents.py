import asyncio
import io
import uuid
from datetime import datetime
from pathlib import Path
from urllib.parse import urlparse

import httpx
import magic
import structlog
from bs4 import BeautifulSoup
from fastapi import APIRouter, Depends, File, HTTPException, Request, UploadFile
from fastapi.responses import JSONResponse
from pydantic import BaseModel, Field
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.api.deps import (
    ResolvedLLMConfig,
    get_current_user,
    get_llm_config,
    require_workspace_member,
)
from app.core.limiter import limiter
from app.core.security import resolve_api_key
from app.db.models import Document, User, UserSettings, WorkspaceMember
from app.db.session import get_db
from app.rag.indexer import index_document
from app.services.document_service import (
    delete_file,
    delete_vectors,
    sync_filename_to_vectors,
    upload_file,
    user_collection_name,
)
from app.tools.web_fetch_tool import is_safe_url

logger = structlog.get_logger(__name__)

router = APIRouter(prefix="/api/documents", tags=["documents"])

_PRIVILEGED_ROLES = frozenset({"owner", "admin"})

ALLOWED_TYPES = {"pdf", "txt", "md", "docx"}
MAX_SIZE = 50 * 1024 * 1024

# MIME type prefixes that are permitted for upload.
# python-magic inspects the actual file bytes, so extension spoofing is caught.
ALLOWED_MIME_PREFIXES = (
    "text/",
    "application/pdf",
    "application/msword",
    "application/vnd.openxmlformats",
    "application/vnd.ms-",
    "image/png",
    "image/jpeg",
    "image/gif",
    "image/webp",
)


async def _assert_doc_write_access(
    doc: "Document",
    user: User,
    db: AsyncSession,
) -> None:
    """Raise HTTPException if user lacks write permission on this document.

    Workspace documents: allow owner/admin members.
    Personal documents: allow only the original uploader.
    """
    if doc.workspace_id is not None:
        membership = await db.scalar(
            select(WorkspaceMember).where(
                WorkspaceMember.workspace_id == doc.workspace_id,
                WorkspaceMember.user_id == user.id,
            )
        )
        if not membership or membership.role not in _PRIVILEGED_ROLES:
            raise HTTPException(status_code=403, detail="Not authorized")
    elif doc.user_id != user.id:
        raise HTTPException(status_code=404, detail="Document not found")


async def _get_qdrant_collection(
    workspace_id: uuid.UUID | None,
    user: User,
    db: AsyncSession,
) -> str:
    """Get the Qdrant collection name for a user or workspace document.

    For workspace documents: validates membership via require_workspace_member
    and returns the workspace collection name.
    For personal documents: returns the user collection name.

    Raises HTTPException 404 if workspace is missing/deleted/outside org,
    or 403 if user is not a member.
    """
    if workspace_id is None:
        return user_collection_name(str(user.id))
    await require_workspace_member(workspace_id, user, db)
    return f"workspace_{workspace_id}"


class DocumentOut(BaseModel):
    id: uuid.UUID
    filename: str
    file_type: str
    source_url: str | None = None
    file_size_bytes: int
    chunk_count: int
    created_at: datetime
    workspace_id: uuid.UUID | None = None
    model_config = {"from_attributes": True}


class DocumentRename(BaseModel):
    filename: str = Field(min_length=1, max_length=255)


def extract_text(content: bytes, file_type: str) -> str:
    match file_type:
        case "txt" | "md":
            return content.decode("utf-8", errors="ignore")
        case "pdf":
            import pypdf

            reader = pypdf.PdfReader(io.BytesIO(content))
            return "\n".join(p.extract_text() or "" for p in reader.pages)
        case "docx":
            import docx

            doc = docx.Document(io.BytesIO(content))
            return "\n".join(p.text for p in doc.paragraphs)
        case _:
            return ""


@router.get("", response_model=dict[str, list[DocumentOut]])
@limiter.limit("60/minute")
async def list_documents(
    request: Request,
    workspace_id: uuid.UUID | None = None,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
) -> dict[str, list[DocumentOut]]:
    if workspace_id is not None:
        # Workspace listing: show all docs in this workspace to members.
        await require_workspace_member(workspace_id, user, db)
        query = (
            select(Document)
            .where(
                Document.workspace_id == workspace_id,
                Document.is_deleted.is_(False),
            )
            .order_by(Document.created_at.desc())
        )
    else:
        # Personal listing: only the user's own non-workspace documents.
        query = (
            select(Document)
            .where(
                Document.user_id == user.id,
                Document.workspace_id.is_(None),
                Document.is_deleted.is_(False),
            )
            .order_by(Document.created_at.desc())
        )
    rows = await db.scalars(query)
    docs = rows.all()
    return {"documents": [DocumentOut.model_validate(d) for d in docs]}


@router.patch("/{doc_id}", response_model=DocumentOut)
@limiter.limit("30/minute")
async def rename_document(
    request: Request,
    doc_id: uuid.UUID,
    body: DocumentRename,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
) -> DocumentOut:
    doc = await db.get(Document, doc_id)
    if not doc or doc.is_deleted:
        raise HTTPException(status_code=404, detail="Document not found")
    await _assert_doc_write_access(doc, user, db)
    doc.filename = body.filename
    await db.commit()

    # Sync new filename to Qdrant vector payload so search results stay fresh.
    collection = doc.qdrant_collection or user_collection_name(str(user.id))
    try:
        await sync_filename_to_vectors(collection, str(doc.id), body.filename)
    except Exception as exc:
        logger.warning("qdrant_rename_sync_failed", doc_id=str(doc.id), error=str(exc))

    return DocumentOut.model_validate(doc)


@router.delete("/{doc_id}", status_code=204)
@limiter.limit("30/minute")
async def delete_document(
    request: Request,
    doc_id: uuid.UUID,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
) -> None:
    doc = await db.get(Document, doc_id)
    if not doc or doc.is_deleted:
        raise HTTPException(status_code=404, detail="Document not found")
    await _assert_doc_write_access(doc, user, db)

    doc.is_deleted = True
    doc.chunk_count = 0
    await db.commit()

    # Use the stored collection name so workspace docs are cleaned up correctly.
    collection = doc.qdrant_collection or user_collection_name(str(user.id))
    try:
        await delete_vectors(collection, str(doc.id))
    except Exception:
        logger.exception("qdrant_delete_failed", doc_id=str(doc.id))
        logger.info("document_deleted", user_id=str(user.id), doc_id=str(doc.id))
        return JSONResponse(
            status_code=207,
            content={
                "status": "partial",
                "message": "Document deleted but vector cleanup failed",
            },
        )

    logger.info("document_deleted", user_id=str(user.id), doc_id=str(doc.id))


@router.post("", status_code=201, response_model=DocumentOut)
@limiter.limit("10/minute")
async def upload_document(
    request: Request,
    file: UploadFile = File(...),
    workspace_id: uuid.UUID | None = None,
    user: User = Depends(get_current_user),
    llm: ResolvedLLMConfig = Depends(get_llm_config),
    db: AsyncSession = Depends(get_db),
) -> DocumentOut:
    ext = (file.filename or "").rsplit(".", 1)[-1].lower()
    if ext not in ALLOWED_TYPES:
        raise HTTPException(status_code=400, detail=f"File type .{ext} not supported")

    content = await file.read()
    if len(content) > MAX_SIZE:
        raise HTTPException(status_code=400, detail="File exceeds 50MB limit")

    # MIME magic-byte validation — catches files with spoofed extensions.
    detected_mime = magic.from_buffer(content[:2048], mime=True)
    if not detected_mime.startswith(ALLOWED_MIME_PREFIXES):
        raise HTTPException(
            status_code=400,
            detail=f"File type not allowed: {detected_mime}",
        )

    # Validate workspace membership before touching MinIO to avoid orphaned objects.
    qdrant_collection = await _get_qdrant_collection(workspace_id, user, db)

    safe_name = Path(file.filename or "upload").name
    object_key = f"{user.id}/{uuid.uuid4()}_{safe_name}"

    try:
        # Run upload and text extraction concurrently — neither depends on the other.
        _, text = await asyncio.gather(
            upload_file(content, object_key),
            asyncio.to_thread(extract_text, content, ext),
        )
        doc = Document(
            user_id=user.id,
            filename=safe_name,
            file_type=ext,
            file_size_bytes=len(content),
            qdrant_collection=qdrant_collection,
            minio_object_key=object_key,
        )
        if workspace_id is not None:
            doc.workspace_id = workspace_id
        db.add(doc)
        await db.flush()
        # Embeddings always use OpenAI (text-embedding-3-small), resolve the
        # OpenAI key regardless of which LLM provider the user has selected.
        openai_key = resolve_api_key("openai", llm.raw_keys)
        if not openai_key:
            raise HTTPException(
                status_code=400,
                detail="OpenAI API key is required for document embedding. "
                "Configure it in Settings or ask the admin.",
            )
        chunk_count = await index_document(
            str(user.id),
            str(doc.id),
            text,
            openai_key,
            doc_name=safe_name,
            collection_name=qdrant_collection,
        )
        doc.chunk_count = chunk_count
        await db.commit()
    except Exception:
        # Best-effort orphan cleanup: remove the MinIO object if anything
        # downstream (DB flush, indexing, etc.) fails after the upload.
        try:
            await delete_file(object_key)
        except Exception:
            logger.warning(
                "minio_orphan_cleanup_failed",
                object_key=object_key,
                user_id=str(user.id),
            )
        raise

    logger.info(
        "document_uploaded",
        user_id=str(user.id),
        doc_id=str(doc.id),
        filename=doc.filename,
        file_type=ext,
        file_size_bytes=len(content),
        chunk_count=chunk_count,
    )
    return DocumentOut.model_validate(doc)


class IngestUrlRequest(BaseModel):
    url: str
    workspace_id: uuid.UUID | None = None


_MAX_URL_CONTENT_BYTES = 5 * 1024 * 1024  # 5 MB


def _extract_page_text(html_bytes: bytes, url: str = "") -> tuple[str, str]:
    """Return (title, body_text) from HTML. Strips noise, collects content tags."""
    soup = BeautifulSoup(html_bytes, "lxml")
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()
    title = ""
    if soup.title and soup.title.string:
        title = soup.title.string.strip()
    if not title:
        h1 = soup.find("h1")
        if h1:
            title = h1.get_text(strip=True)
    parts = []
    for tag in soup.find_all(["p", "h1", "h2", "h3", "h4", "h5", "h6", "li", "td"]):
        text = tag.get_text(separator=" ", strip=True)
        if text:
            parts.append(text)
    fallback = urlparse(url).hostname or "Web Page"
    return title or fallback, "\n\n".join(parts)


@router.post("/ingest-url", response_model=DocumentOut, status_code=201)
@limiter.limit("10/minute")
async def ingest_url(  # noqa: C901
    request: Request,
    body: IngestUrlRequest,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
) -> DocumentOut:
    """Fetch a web page and add it to the knowledge base."""
    if not await is_safe_url(body.url):
        raise HTTPException(
            status_code=400,
            detail=f"URL not allowed (internal or non-http): {body.url!r}",
        )

    workspace_id = body.workspace_id
    qdrant_collection = await _get_qdrant_collection(workspace_id, user, db)

    # Resolve OpenAI key before fetching the URL (fail fast).
    user_settings = await db.scalar(
        select(UserSettings).where(UserSettings.user_id == user.id)
    )
    raw_keys = user_settings.api_keys if user_settings else {}
    openai_key = resolve_api_key("openai", raw_keys)
    if not openai_key:
        raise HTTPException(
            status_code=400,
            detail=(
                "OpenAI API key required for document embedding. "
                "Configure it in Settings."
            ),
        )

    async with httpx.AsyncClient(follow_redirects=True, timeout=15.0) as client:
        try:
            response = await client.get(body.url)
            response.raise_for_status()
        except httpx.HTTPError as e:
            raise HTTPException(
                status_code=400, detail=f"Failed to fetch URL: {e}"
            ) from e
        content_type = response.headers.get("content-type", "")
        if not content_type.startswith("text/"):
            raise HTTPException(
                status_code=400,
                detail=(
                    f"URL returned non-text content ({content_type!r}). "
                    "Only HTML and plain-text pages are supported."
                ),
            )
        content_length = response.headers.get("content-length")
        if content_length and int(content_length) > _MAX_URL_CONTENT_BYTES:
            raise HTTPException(status_code=400, detail="Page too large (max 5 MB)")
        if len(response.content) > _MAX_URL_CONTENT_BYTES:
            raise HTTPException(status_code=400, detail="Page too large (max 5 MB)")
        html_content = response.content

    title, text = await asyncio.to_thread(_extract_page_text, html_content, body.url)
    if not text.strip():
        raise HTTPException(status_code=400, detail="No readable content found on page")

    text_bytes = text.encode("utf-8")
    object_key = f"{user.id}/{uuid.uuid4()}_webpage.txt"

    try:
        await upload_file(text_bytes, object_key)

        doc = Document(
            user_id=user.id,
            filename=title[:255],
            file_type="txt",
            file_size_bytes=len(text_bytes),
            qdrant_collection=qdrant_collection,
            minio_object_key=object_key,
            source_url=body.url,
        )
        if workspace_id is not None:
            doc.workspace_id = workspace_id
        db.add(doc)
        await db.flush()

        chunk_count = await index_document(
            str(user.id),
            str(doc.id),
            text,
            openai_key,
            doc_name=title,
            collection_name=qdrant_collection,
        )
        doc.chunk_count = chunk_count
        await db.commit()
    except Exception:
        try:
            await delete_file(object_key)
        except Exception:
            logger.warning(
                "minio_orphan_cleanup_failed",
                object_key=object_key,
                user_id=str(user.id),
            )
        raise
    logger.info(
        "document_url_ingested",
        user_id=str(user.id),
        doc_id=str(doc.id),
        url=body.url,
    )
    return DocumentOut.model_validate(doc)
